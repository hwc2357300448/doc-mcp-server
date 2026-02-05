"""内容编辑工具"""
import os
from typing import Optional, Dict, Any, List
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..utils import DocumentManager, validate_file_path, handle_docx_errors

# 全局文档管理器实例
doc_manager = DocumentManager()


@handle_docx_errors
async def add_paragraph(
    filename: str,
    text: str,
    style: Optional[str] = None,
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: bool = False,
    italic: bool = False,
    color: Optional[str] = None,
    highlight: Optional[str] = None,
    first_line_indent: Optional[float] = None,
    left_indent: Optional[float] = None,
    right_indent: Optional[float] = None,
    alignment: Optional[str] = None
) -> Dict[str, Any]:
    """
    添加段落到Word文档

    参数:
        filename: 文档路径
        text: 段落文本内容
        style: 段落样式名称（可选）
        font_name: 字体名称，如'宋体'、'Arial'（可选）
        font_size: 字号，单位磅（可选）
        bold: 是否粗体（默认False）
        italic: 是否斜体（默认False）
        color: 文字颜色，十六进制RGB，如'FF0000'（可选）
        highlight: 背景色（高亮），十六进制RGB，如'FFFF00'（可选）
        first_line_indent: 首行缩进，单位厘米（可选）
        left_indent: 左缩进，单位厘米（可选）
        right_indent: 右缩进，单位厘米（可选）
        alignment: 对齐方式，可选值：'left'、'center'、'right'、'justify'（可选）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    # 添加段落
    para = doc.add_paragraph(text, style=style)

    # 设置字体格式
    if any([font_name, font_size, bold, italic, color, highlight]):
        for run in para.runs:
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = Pt(font_size)
            if bold is not None:
                run.font.bold = bold
            if italic is not None:
                run.font.italic = italic
            if color:
                run.font.color.rgb = RGBColor(
                    int(color[0:2], 16),
                    int(color[2:4], 16),
                    int(color[4:6], 16)
                )
            if highlight:
                # 设置背景色（高亮）
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), highlight)
                run._element.get_or_add_rPr().append(shd)

    # 设置段落格式：缩进
    if first_line_indent is not None:
        para.paragraph_format.first_line_indent = Inches(first_line_indent / 2.54)
    if left_indent is not None:
        para.paragraph_format.left_indent = Inches(left_indent / 2.54)
    if right_indent is not None:
        para.paragraph_format.right_indent = Inches(right_indent / 2.54)

    # 设置段落格式：对齐方式
    if alignment:
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if alignment.lower() in alignment_map:
            para.paragraph_format.alignment = alignment_map[alignment.lower()]

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": "段落添加成功",
        "paragraph_index": len(doc.paragraphs) - 1
    }


@handle_docx_errors
async def batch_add_paragraphs(
    filename: str,
    paragraphs: List[Dict[str, Any]]
) -> Dict[str, Any]:
    """
    批量添加多个段落到Word文档

    参数:
        filename: 文档路径
        paragraphs: 段落列表，每个段落是一个字典，包含：
            - text: 段落文本（必需）
            - style: 段落样式（可选）
            - font_name: 字体名称（可选）
            - font_size: 字号（可选）
            - bold: 是否粗体（可选）
            - italic: 是否斜体（可选）
            - color: 文字颜色（可选）
            - highlight: 背景色（高亮）（可选）
            - first_line_indent: 首行缩进，单位厘米（可选）
            - left_indent: 左缩进，单位厘米（可选）
            - right_indent: 右缩进，单位厘米（可选）
            - alignment: 对齐方式（可选）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if not paragraphs:
        raise ValueError("段落列表不能为空")

    added_count = 0
    for para_data in paragraphs:
        text = para_data.get('text')
        if not text:
            continue

        style = para_data.get('style')
        font_name = para_data.get('font_name')
        font_size = para_data.get('font_size')
        bold = para_data.get('bold', False)
        italic = para_data.get('italic', False)
        color = para_data.get('color')
        highlight = para_data.get('highlight')
        first_line_indent = para_data.get('first_line_indent')
        left_indent = para_data.get('left_indent')
        right_indent = para_data.get('right_indent')
        alignment = para_data.get('alignment')

        # 添加段落
        para = doc.add_paragraph(text, style=style)

        # 设置字体格式
        if any([font_name, font_size, bold, italic, color, highlight]):
            for run in para.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold:
                    run.font.bold = bold
                if italic:
                    run.font.italic = italic
                if color:
                    run.font.color.rgb = RGBColor(
                        int(color[0:2], 16),
                        int(color[2:4], 16),
                        int(color[4:6], 16)
                    )
                if highlight:
                    # 设置背景色（高亮）
                    from docx.oxml.shared import OxmlElement
                    from docx.oxml.ns import qn
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), highlight)
                    run._element.get_or_add_rPr().append(shd)

        # 设置段落格式：缩进
        if first_line_indent is not None:
            para.paragraph_format.first_line_indent = Inches(first_line_indent / 2.54)
        if left_indent is not None:
            para.paragraph_format.left_indent = Inches(left_indent / 2.54)
        if right_indent is not None:
            para.paragraph_format.right_indent = Inches(right_indent / 2.54)

        # 设置段落格式：对齐方式
        if alignment:
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if alignment.lower() in alignment_map:
                para.paragraph_format.alignment = alignment_map[alignment.lower()]

        added_count += 1

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"批量添加成功，共添加{added_count}个段落",
        "added_count": added_count
    }


@handle_docx_errors
async def add_heading(
    filename: str,
    text: str,
    level: int = 1,
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: bool = False,
    italic: bool = False,
    color: Optional[str] = None,
    highlight: Optional[str] = None,
    first_line_indent: Optional[float] = None,
    left_indent: Optional[float] = None,
    right_indent: Optional[float] = None,
    alignment: Optional[str] = None
) -> Dict[str, Any]:
    """
    添加标题到Word文档

    参数:
        filename: 文档路径
        text: 标题文本
        level: 标题级别（1-9，默认1）
        font_name: 字体名称（可选）
        font_size: 字号（可选）
        bold: 是否粗体（默认False）
        italic: 是否斜体（默认False）
        color: 文字颜色，十六进制RGB（可选）
        highlight: 背景色（高亮），十六进制RGB（可选）
        first_line_indent: 首行缩进，单位厘米（可选）
        left_indent: 左缩进，单位厘米（可选）
        right_indent: 右缩进，单位厘米（可选）
        alignment: 对齐方式（可选）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if not 1 <= level <= 9:
        raise ValueError(f"标题级别必须在1-9之间，当前值: {level}")

    # 添加标题
    heading = doc.add_heading(text, level=level)

    # 设置字体格式
    if any([font_name, font_size, bold, italic, color, highlight]):
        for run in heading.runs:
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = Pt(font_size)
            if bold is not None:
                run.font.bold = bold
            if italic is not None:
                run.font.italic = italic
            if color:
                run.font.color.rgb = RGBColor(
                    int(color[0:2], 16),
                    int(color[2:4], 16),
                    int(color[4:6], 16)
                )
            if highlight:
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), highlight)
                run._element.get_or_add_rPr().append(shd)

    # 设置段落格式：缩进
    if first_line_indent is not None:
        heading.paragraph_format.first_line_indent = Inches(first_line_indent / 2.54)
    if left_indent is not None:
        heading.paragraph_format.left_indent = Inches(left_indent / 2.54)
    if right_indent is not None:
        heading.paragraph_format.right_indent = Inches(right_indent / 2.54)

    # 设置段落格式：对齐方式
    if alignment:
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if alignment.lower() in alignment_map:
            heading.paragraph_format.alignment = alignment_map[alignment.lower()]

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"标题添加成功（级别{level}）",
        "heading_level": level
    }


@handle_docx_errors
async def delete_paragraph(filename: str, paragraph_index: int) -> Dict[str, Any]:
    """
    删除指定段落

    参数:
        filename: 文档路径
        paragraph_index: 段落索引（从0开始）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise ValueError(f"段落索引超出范围: {paragraph_index}，文档共有{len(doc.paragraphs)}个段落")

    # 删除段落
    para = doc.paragraphs[paragraph_index]
    p_element = para._element
    p_element.getparent().remove(p_element)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"段落{paragraph_index}删除成功"
    }


@handle_docx_errors
async def insert_paragraph(
    filename: str,
    text: str,
    position: int,
    style: Optional[str] = None,
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: bool = False,
    italic: bool = False,
    color: Optional[str] = None,
    highlight: Optional[str] = None,
    first_line_indent: Optional[float] = None,
    left_indent: Optional[float] = None,
    right_indent: Optional[float] = None,
    alignment: Optional[str] = None
) -> Dict[str, Any]:
    """
    在指定位置插入段落

    参数:
        filename: 文档路径
        text: 段落文本内容
        position: 插入位置索引（从0开始，0表示插入到开头）
        style: 段落样式名称（可选）
        font_name: 字体名称（可选）
        font_size: 字号（可选）
        bold: 是否粗体（默认False）
        italic: 是否斜体（默认False）
        color: 文字颜色，十六进制RGB（可选）
        highlight: 背景色（高亮），十六进制RGB（可选）
        first_line_indent: 首行缩进，单位厘米（可选）
        left_indent: 左缩进，单位厘米（可选）
        right_indent: 右缩进，单位厘米（可选）
        alignment: 对齐方式（可选）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if position < 0 or position > len(doc.paragraphs):
        raise ValueError(f"插入位置超出范围: {position}，有效范围: 0-{len(doc.paragraphs)}")

    # 在指定位置插入段落
    if position == len(doc.paragraphs):
        # 在末尾插入
        para = doc.add_paragraph(text, style=style)
    else:
        # 在中间位置插入
        target_para = doc.paragraphs[position]
        new_para = target_para.insert_paragraph_before(text, style=style)
        para = new_para

    # 设置字体格式
    if any([font_name, font_size, bold, italic, color, highlight]):
        for run in para.runs:
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = Pt(font_size)
            if bold:
                run.font.bold = bold
            if italic:
                run.font.italic = italic
            if color:
                run.font.color.rgb = RGBColor(
                    int(color[0:2], 16),
                    int(color[2:4], 16),
                    int(color[4:6], 16)
                )
            if highlight:
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), highlight)
                run._element.get_or_add_rPr().append(shd)

    # 设置段落格式：缩进
    if first_line_indent is not None:
        para.paragraph_format.first_line_indent = Inches(first_line_indent / 2.54)
    if left_indent is not None:
        para.paragraph_format.left_indent = Inches(left_indent / 2.54)
    if right_indent is not None:
        para.paragraph_format.right_indent = Inches(right_indent / 2.54)

    # 设置段落格式：对齐方式
    if alignment:
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if alignment.lower() in alignment_map:
            para.paragraph_format.alignment = alignment_map[alignment.lower()]

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"段落已插入到位置{position}",
        "position": position
    }


@handle_docx_errors
async def delete_paragraph_range(
    filename: str,
    start_index: int,
    end_index: int
) -> Dict[str, Any]:
    """
    删除指定范围的段落

    参数:
        filename: 文档路径
        start_index: 起始段落索引（从0开始，包含）
        end_index: 结束段落索引（从0开始，包含）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    total_paragraphs = len(doc.paragraphs)

    if start_index < 0 or end_index >= total_paragraphs:
        raise ValueError(f"段落索引超出范围，文档共有{total_paragraphs}个段落")

    if start_index > end_index:
        raise ValueError(f"起始索引({start_index})不能大于结束索引({end_index})")

    # 从后往前删除，避免索引变化问题
    deleted_count = 0
    for i in range(end_index, start_index - 1, -1):
        para = doc.paragraphs[i]
        p_element = para._element
        p_element.getparent().remove(p_element)
        deleted_count += 1

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"已删除段落{start_index}到{end_index}，共{deleted_count}个段落",
        "deleted_count": deleted_count
    }


@handle_docx_errors
async def replace_paragraph_range(
    filename: str,
    start_index: int,
    end_index: int,
    new_text: str,
    style: Optional[str] = None,
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: bool = False,
    italic: bool = False,
    color: Optional[str] = None,
    highlight: Optional[str] = None,
    first_line_indent: Optional[float] = None,
    left_indent: Optional[float] = None,
    right_indent: Optional[float] = None,
    alignment: Optional[str] = None
) -> Dict[str, Any]:
    """
    替换指定范围的段落为新内容

    参数:
        filename: 文档路径
        start_index: 起始段落索引（从0开始，包含）
        end_index: 结束段落索引（从0开始，包含）
        new_text: 新的段落文本内容
        style: 段落样式名称（可选）
        font_name: 字体名称（可选）
        font_size: 字号（可选）
        bold: 是否粗体（默认False）
        italic: 是否斜体（默认False）
        color: 文字颜色，十六进制RGB（可选）
        highlight: 背景色（高亮），十六进制RGB（可选）
        first_line_indent: 首行缩进，单位厘米（可选）
        left_indent: 左缩进，单位厘米（可选）
        right_indent: 右缩进，单位厘米（可选）
        alignment: 对齐方式（可选）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    total_paragraphs = len(doc.paragraphs)

    if start_index < 0 or end_index >= total_paragraphs:
        raise ValueError(f"段落索引超出范围，文档共有{total_paragraphs}个段落")

    if start_index > end_index:
        raise ValueError(f"起始索引({start_index})不能大于结束索引({end_index})")

    # 先删除范围内的段落（从后往前删除）
    for i in range(end_index, start_index, -1):
        para = doc.paragraphs[i]
        p_element = para._element
        p_element.getparent().remove(p_element)

    # 在起始位置插入新段落
    if start_index < len(doc.paragraphs):
        target_para = doc.paragraphs[start_index]
        new_para = target_para.insert_paragraph_before(new_text, style=style)
        # 删除原来的起始段落
        p_element = target_para._element
        p_element.getparent().remove(p_element)
    else:
        # 如果是最后一个位置，直接添加
        new_para = doc.add_paragraph(new_text, style=style)

    # 设置字体格式
    if any([font_name, font_size, bold, italic, color, highlight]):
        for run in new_para.runs:
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = Pt(font_size)
            if bold:
                run.font.bold = bold
            if italic:
                run.font.italic = italic
            if color:
                run.font.color.rgb = RGBColor(
                    int(color[0:2], 16),
                    int(color[2:4], 16),
                    int(color[4:6], 16)
                )
            if highlight:
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), highlight)
                run._element.get_or_add_rPr().append(shd)

    # 设置段落格式：缩进
    if first_line_indent is not None:
        new_para.paragraph_format.first_line_indent = Inches(first_line_indent / 2.54)
    if left_indent is not None:
        new_para.paragraph_format.left_indent = Inches(left_indent / 2.54)
    if right_indent is not None:
        new_para.paragraph_format.right_indent = Inches(right_indent / 2.54)

    # 设置段落格式：对齐方式
    if alignment:
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if alignment.lower() in alignment_map:
            new_para.paragraph_format.alignment = alignment_map[alignment.lower()]

    doc_manager.save(abs_path, doc)

    replaced_count = end_index - start_index + 1
    return {
        "success": True,
        "message": f"已替换段落{start_index}到{end_index}，共{replaced_count}个段落",
        "replaced_count": replaced_count
    }


@handle_docx_errors
async def find_text(
    filename: str,
    text_to_find: str,
    match_case: bool = True,
    whole_word: bool = False
) -> Dict[str, Any]:
    """
    在文档中查找文本

    参数:
        filename: 文档路径
        text_to_find: 要查找的文本
        match_case: 是否区分大小写（默认True）
        whole_word: 是否全字匹配（默认False）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    occurrences = []
    search_text = text_to_find if match_case else text_to_find.lower()

    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text if match_case else para.text.lower()

        if whole_word:
            import re
            pattern = r'\b' + re.escape(search_text) + r'\b'
            matches = list(re.finditer(pattern, para_text))
            if matches:
                occurrences.append({
                    "paragraph_index": para_idx,
                    "text": para.text,
                    "count": len(matches)
                })
        else:
            if search_text in para_text:
                count = para_text.count(search_text)
                occurrences.append({
                    "paragraph_index": para_idx,
                    "text": para.text,
                    "count": count
                })

    return {
        "success": True,
        "search_text": text_to_find,
        "total_occurrences": sum(occ["count"] for occ in occurrences),
        "paragraphs_found": len(occurrences),
        "occurrences": occurrences
    }


@handle_docx_errors
async def replace_text(
    filename: str,
    find_text: str,
    replace_text: str
) -> Dict[str, Any]:
    """
    替换文档中的文本

    参数:
        filename: 文档路径
        find_text: 要查找的文本
        replace_text: 替换后的文本
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    replacement_count = 0

    # 遍历所有段落
    for para in doc.paragraphs:
        if find_text in para.text:
            # 替换段落中的文本
            inline = para.runs
            for run in inline:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    replacement_count += 1

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"替换完成，共替换{replacement_count}处",
        "find_text": find_text,
        "replace_text": replace_text,
        "replacement_count": replacement_count
    }

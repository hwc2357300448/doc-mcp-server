"""
接口文档生成工具
功能: 在指定位置插入标准格式的接口文档表格
"""

from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from ..utils import DocumentManager, validate_file_path, handle_docx_errors

# 全局文档管理器实例
doc_manager = DocumentManager()


def set_cell_background(cell, color_hex: str):
    """
    设置单元格背景色

    参数:
        cell: 单元格对象
        color_hex: 十六进制颜色值，如 'D9D9D9'
    """
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.find(qn('w:shd'))
    if shading is None:
        shading = shading_elm.makeelement(qn('w:shd'), {})
        shading_elm.append(shading)
    shading.set(qn('w:fill'), color_hex)


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 10.5):
    """
    设置单元格文本及格式

    参数:
        cell: 单元格对象
        text: 文本内容
        bold: 是否粗体
        font_size: 字号
    """
    # 清空单元格
    cell.text = ''

    # 添加段落
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 添加文本
    run = paragraph.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    run.font.bold = bold


@handle_docx_errors
async def insert_interface_doc(
    filename: str,
    position: int,
    name: str,
    path: str,
    description: str,
    method: str = 'POST',
    request_params: Optional[List[List[str]]] = None,
    response_params: Optional[List[List[str]]] = None,
    request_example: Optional[str] = None,
    response_example: Optional[str] = None
) -> Dict[str, Any]:
    """
    在指定位置插入接口文档表格

    参数:
        filename: 文档路径
        position: 插入位置（段落索引，0表示开头）
        name: 接口名称
        path: 访问路径
        description: 服务说明
        method: 请求方式（默认POST）
        request_params: 请求参数列表，每个参数为 [字段名, 类型, 是否必填, 说明]
        response_params: 响应参数列表，每个参数为 [字段名, 类型, 是否必填, 说明]
        request_example: 请求示例（可选）
        response_example: 响应示例（可选）

    返回:
        包含成功状态和消息的字典
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if request_params is None:
        request_params = []
    if response_params is None:
        response_params = []

    # 计算表格行数
    rows = 4  # 服务名称、访问路径、请求方式、服务说明
    rows += 2  # 请求参数标题行 + 字段表头
    rows += len(request_params)
    rows += 2  # 响应参数标题行 + 字段表头
    rows += len(response_params)

    if request_example:
        rows += 2  # 请求示例标题行 + 内容行
    if response_example:
        rows += 2  # 响应示例标题行 + 内容行

    # 在指定位置插入表格
    if position >= len(doc.paragraphs):
        # 如果位置超出范围，在末尾添加
        table = doc.add_table(rows=rows, cols=4)
    else:
        # 在指定段落前插入表格
        target_para = doc.paragraphs[position]
        table = doc.add_table(rows=rows, cols=4)
        # 将表格移动到目标位置
        table_element = table._element
        target_para._element.addprevious(table_element)

    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    current_row = 0

    # 第1行: 服务名称
    cell = table.rows[current_row].cells[0]
    set_cell_text(cell, '服务名称', bold=False)
    set_cell_background(cell, 'D9D9D9')
    merged_cell = table.rows[current_row].cells[1].merge(table.rows[current_row].cells[3])
    set_cell_text(merged_cell, name, bold=False)
    current_row += 1

    # 第2行: 访问路径
    cell = table.rows[current_row].cells[0]
    set_cell_text(cell, '访问路径', bold=False)
    set_cell_background(cell, 'D9D9D9')
    merged_cell = table.rows[current_row].cells[1].merge(table.rows[current_row].cells[3])
    set_cell_text(merged_cell, path, bold=False)
    current_row += 1

    # 第3行: 请求方式
    cell = table.rows[current_row].cells[0]
    set_cell_text(cell, '请求方式', bold=False)
    set_cell_background(cell, 'D9D9D9')
    merged_cell = table.rows[current_row].cells[1].merge(table.rows[current_row].cells[3])
    set_cell_text(merged_cell, method, bold=False)
    current_row += 1

    # 第4行: 服务说明
    cell = table.rows[current_row].cells[0]
    set_cell_text(cell, '服务说明', bold=False)
    set_cell_background(cell, 'D9D9D9')
    merged_cell = table.rows[current_row].cells[1].merge(table.rows[current_row].cells[3])
    set_cell_text(merged_cell, description, bold=False)
    current_row += 1

    # 请求参数标题行
    for col_idx in range(4):
        set_cell_background(table.rows[current_row].cells[col_idx], 'D9D9D9')
    merged_cell = table.rows[current_row].cells[0].merge(table.rows[current_row].cells[3])
    set_cell_text(merged_cell, '请求参数', bold=False)
    current_row += 1

    # 请求参数表头
    headers = ['字段名', '类型', '是否必填', '说明']
    for col_idx, header in enumerate(headers):
        cell = table.rows[current_row].cells[col_idx]
        set_cell_text(cell, header, bold=False)
        set_cell_background(cell, 'D9D9D9')
    current_row += 1

    # 请求参数列表
    for param in request_params:
        for col_idx, value in enumerate(param):
            cell = table.rows[current_row].cells[col_idx]
            set_cell_text(cell, str(value), bold=False)
        current_row += 1

    # 响应参数标题行
    for col_idx in range(4):
        set_cell_background(table.rows[current_row].cells[col_idx], 'D9D9D9')
    merged_cell = table.rows[current_row].cells[0].merge(table.rows[current_row].cells[3])
    set_cell_text(merged_cell, '响应参数', bold=False)
    current_row += 1

    # 响应参数表头
    for col_idx, header in enumerate(headers):
        cell = table.rows[current_row].cells[col_idx]
        set_cell_text(cell, header, bold=False)
        set_cell_background(cell, 'D9D9D9')
    current_row += 1

    # 响应参数列表
    for param in response_params:
        for col_idx, value in enumerate(param):
            cell = table.rows[current_row].cells[col_idx]
            set_cell_text(cell, str(value), bold=False)
        current_row += 1

    # 请求示例（可选）
    if request_example:
        for col_idx in range(4):
            set_cell_background(table.rows[current_row].cells[col_idx], 'D9D9D9')
        merged_cell = table.rows[current_row].cells[0].merge(table.rows[current_row].cells[3])
        set_cell_text(merged_cell, '请求示例', bold=False)
        current_row += 1

        merged_cell = table.rows[current_row].cells[0].merge(table.rows[current_row].cells[3])
        set_cell_text(merged_cell, request_example, bold=False)
        current_row += 1

    # 响应示例（可选）
    if response_example:
        for col_idx in range(4):
            set_cell_background(table.rows[current_row].cells[col_idx], 'D9D9D9')
        merged_cell = table.rows[current_row].cells[0].merge(table.rows[current_row].cells[3])
        set_cell_text(merged_cell, '响应示例', bold=False)
        current_row += 1

        merged_cell = table.rows[current_row].cells[0].merge(table.rows[current_row].cells[3])
        set_cell_text(merged_cell, response_example, bold=False)
        current_row += 1

    # 保存文档
    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"接口文档表格已插入到位置{position}",
        "interface_name": name,
        "position": position
    }

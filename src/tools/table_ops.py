"""表格操作工具"""
import os
from typing import Optional, Dict, Any, List
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from ..utils import DocumentManager, validate_file_path, handle_docx_errors

# 全局文档管理器实例
doc_manager = DocumentManager()


@handle_docx_errors
async def add_table(
    filename: str,
    rows: int,
    cols: int,
    data: Optional[List[List[str]]] = None
) -> Dict[str, Any]:
    """
    创建表格

    参数:
        filename: 文档路径
        rows: 行数
        cols: 列数
        data: 表格数据（可选），二维列表
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if rows <= 0 or cols <= 0:
        raise ValueError(f"行数和列数必须大于0，当前值: rows={rows}, cols={cols}")

    # 创建表格
    table = doc.add_table(rows=rows, cols=cols)

    # 尝试设置表格样式，如果样式不存在则跳过
    try:
        table.style = 'Table Grid'
    except KeyError:
        # 样式不存在，使用默认样式
        pass

    # 填充数据
    if data:
        for i, row_data in enumerate(data):
            if i >= rows:
                break
            for j, cell_data in enumerate(row_data):
                if j >= cols:
                    break
                table.rows[i].cells[j].text = str(cell_data)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"表格创建成功（{rows}行 x {cols}列）",
        "table_index": len(doc.tables) - 1,
        "rows": rows,
        "cols": cols
    }


@handle_docx_errors
async def insert_table(
    filename: str,
    position: int,
    rows: int,
    cols: int,
    data: Optional[List[List[str]]] = None
) -> Dict[str, Any]:
    """
    在指定位置插入表格（在指定索引之后插入）

    参数:
        filename: 文档路径
        position: 插入位置索引（从0开始）
                 表格将插入到指定索引之后
                 例如：position=0 表示插入到索引0之后
                      position=5 表示插入到索引5之后
        rows: 行数
        cols: 列数
        data: 表格数据（可选），二维列表
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if position < 0 or position >= len(doc.paragraphs):
        raise ValueError(f"插入位置超出范围: {position}，有效范围: 0-{len(doc.paragraphs)-1}")

    if rows <= 0 or cols <= 0:
        raise ValueError(f"行数和列数必须大于0，当前值: rows={rows}, cols={cols}")

    # 在指定索引之后插入表格
    # 策略：先创建表格，然后移动到正确位置
    table = doc.add_table(rows=rows, cols=cols)

    # 如果不是追加到末尾，需要移动表格位置
    if position + 1 < len(doc.paragraphs):
        # 获取表格元素和目标段落元素
        table_element = table._element
        target_para = doc.paragraphs[position + 1]
        target_element = target_para._element

        # 将表格元素移动到目标段落之前
        target_element.addprevious(table_element)

    # 尝试设置表格样式
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass

    # 填充数据
    if data:
        for i, row_data in enumerate(data):
            if i >= rows:
                break
            for j, cell_data in enumerate(row_data):
                if j >= cols:
                    break
                table.rows[i].cells[j].text = str(cell_data)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"表格已插入到位置{position}之后（{rows}行 x {cols}列）",
        "position": position,
        "table_index": len(doc.tables) - 1,
        "rows": rows,
        "cols": cols
    }


@handle_docx_errors
async def set_table_cell_content(
    filename: str,
    table_index: int,
    row_index: int,
    col_index: int,
    text: str,
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
    highlight: Optional[str] = None,
    alignment: Optional[str] = None
) -> Dict[str, Any]:
    """
    设置表格单元格内容

    参数:
        filename: 文档路径
        table_index: 表格索引（从0开始）
        row_index: 行索引（从0开始）
        col_index: 列索引（从0开始）
        text: 单元格文本内容
        font_name: 字体名称（可选）
        font_size: 字号（可选）
        bold: 是否粗体（可选）
        italic: 是否斜体（可选）
        color: 文字颜色，十六进制RGB（可选）
        highlight: 背景色（高亮），十六进制RGB（可选）
        alignment: 对齐方式（可选）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}，文档共有{len(doc.tables)}个表格")

    table = doc.tables[table_index]

    if row_index < 0 or row_index >= len(table.rows):
        raise ValueError(f"行索引超出范围: {row_index}，表格共有{len(table.rows)}行")

    if col_index < 0 or col_index >= len(table.columns):
        raise ValueError(f"列索引超出范围: {col_index}，表格共有{len(table.columns)}列")

    cell = table.rows[row_index].cells[col_index]
    cell.text = text

    # 设置字体格式
    if any([font_name, font_size, bold, italic, color, highlight]):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                if color:
                    run.font.color.rgb = RGBColor(
                        int(color[0:2], 16),
                        int(color[2:4], 16),
                        int(color[4:6], 16)
                    )
                if highlight:
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), highlight)
                    run._element.get_or_add_rPr().append(shd)

    # 设置段落对齐方式
    if alignment:
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if alignment.lower() in alignment_map:
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignment_map[alignment.lower()]

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"单元格内容设置成功（表格{table_index}，行{row_index}，列{col_index}）"
    }


@handle_docx_errors
async def format_table(
    filename: str,
    table_index: int,
    border_style: Optional[str] = None,
    has_header_row: bool = False
) -> Dict[str, Any]:
    """
    格式化表格

    参数:
        filename: 文档路径
        table_index: 表格索引（从0开始）
        border_style: 边框样式（可选）
        has_header_row: 是否有表头行（默认False）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]

    # 设置表格样式
    if border_style:
        table.style = border_style

    # 如果有表头行，设置第一行为粗体
    if has_header_row and len(table.rows) > 0:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"表格{table_index}格式化成功"
    }


@handle_docx_errors
async def insert_table_row(
    filename: str,
    table_index: int,
    row_index: int
) -> Dict[str, Any]:
    """
    在表格中插入新行

    参数:
        filename: 文档路径
        table_index: 表格索引（从0开始）
        row_index: 插入位置的行索引（从0开始）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]

    if row_index < 0 or row_index > len(table.rows):
        raise ValueError(f"行索引超出范围: {row_index}")

    # 插入新行
    table.add_row()

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"表格{table_index}插入新行成功",
        "new_row_count": len(table.rows)
    }


@handle_docx_errors
async def set_column_width(
    filename: str,
    table_index: int,
    col_index: int,
    width: float
) -> Dict[str, Any]:
    """
    设置表格列宽

    参数:
        filename: 文档路径
        table_index: 表格索引（从0开始）
        col_index: 列索引（从0开始）
        width: 列宽（英寸）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]

    if col_index < 0 or col_index >= len(table.columns):
        raise ValueError(f"列索引超出范围: {col_index}")

    # 设置列宽
    for cell in table.columns[col_index].cells:
        cell.width = Inches(width)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"列宽设置成功（表格{table_index}，列{col_index}，宽度{width}英寸）"
    }


@handle_docx_errors
async def delete_table_row(
    filename: str,
    table_index: int,
    row_index: int
) -> Dict[str, Any]:
    """
    删除表格中的指定行

    参数:
        filename: 文档路径
        table_index: 表格索引（从0开始）
        row_index: 要删除的行索引（从0开始）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]

    if row_index < 0 or row_index >= len(table.rows):
        raise ValueError(f"行索引超出范围: {row_index}")

    # 删除行（通过XML操作）
    row = table.rows[row_index]
    row._element.getparent().remove(row._element)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"表格{table_index}的第{row_index}行删除成功",
        "remaining_rows": len(table.rows)
    }


@handle_docx_errors
async def delete_table_column(
    filename: str,
    table_index: int,
    col_index: int
) -> Dict[str, Any]:
    """
    删除表格中的指定列

    参数:
        filename: 文档路径
        table_index: 表格索引（从0开始）
        col_index: 要删除的列索引（从0开始）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]

    if col_index < 0 or col_index >= len(table.columns):
        raise ValueError(f"列索引超出范围: {col_index}")

    # 删除每一行中的指定列单元格
    for row in table.rows:
        cell = row.cells[col_index]
        cell._element.getparent().remove(cell._element)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"表格{table_index}的第{col_index}列删除成功"
    }


@handle_docx_errors
async def delete_table(
    filename: str,
    table_index: int
) -> Dict[str, Any]:
    """
    删除整个表格

    参数:
        filename: 文档路径
        table_index: 表格索引（从0开始）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]
    table._element.getparent().remove(table._element)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"表格{table_index}删除成功",
        "remaining_tables": len(doc.tables)
    }


@handle_docx_errors
async def merge_table_cells(
    filename: str,
    table_index: int,
    start_row: int,
    start_col: int,
    end_row: int,
    end_col: int
) -> Dict[str, Any]:
    """
    合并表格单元格（矩形区域）

    参数:
        filename: 文档路径
        table_index: 表格索引（从0开始）
        start_row: 起始行索引
        start_col: 起始列索引
        end_row: 结束行索引
        end_col: 结束列索引
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]

    if start_row < 0 or end_row >= len(table.rows) or start_row > end_row:
        raise ValueError(f"行索引无效: start_row={start_row}, end_row={end_row}")

    if start_col < 0 or end_col >= len(table.columns) or start_col > end_col:
        raise ValueError(f"列索引无效: start_col={start_col}, end_col={end_col}")

    # 获取起始单元格
    start_cell = table.cell(start_row, start_col)
    end_cell = table.cell(end_row, end_col)

    # 合并单元格
    start_cell.merge(end_cell)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"单元格合并成功（行{start_row}-{end_row}，列{start_col}-{end_col}）"
    }


@handle_docx_errors
async def set_cell_alignment(
    filename: str,
    table_index: int,
    row_index: int,
    col_index: int,
    horizontal: str = "left",
    vertical: str = "top"
) -> Dict[str, Any]:
    """
    设置单元格对齐方式

    参数:
        filename: 文档路径
        table_index: 表格索引
        row_index: 行索引
        col_index: 列索引
        horizontal: 水平对齐（left/center/right/justify）
        vertical: 垂直对齐（top/center/bottom）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]
    cell = table.cell(row_index, col_index)

    # 设置水平对齐
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    if horizontal in align_map:
        for paragraph in cell.paragraphs:
            paragraph.alignment = align_map[horizontal]

    # 设置垂直对齐
    valign_map = {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    }
    if vertical in valign_map:
        cell.vertical_alignment = valign_map[vertical]

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"单元格对齐设置成功（水平:{horizontal}，垂直:{vertical}）"
    }


@handle_docx_errors
async def set_cell_background(
    filename: str,
    table_index: int,
    row_index: int,
    col_index: int,
    color: str
) -> Dict[str, Any]:
    """
    设置单元格背景色

    参数:
        filename: 文档路径
        table_index: 表格索引
        row_index: 行索引
        col_index: 列索引
        color: 颜色（十六进制，如 'FF0000' 表示红色）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]
    cell = table.cell(row_index, col_index)

    # 设置背景色
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"单元格背景色设置成功（颜色:{color}）"
    }


@handle_docx_errors
async def set_cell_padding(
    filename: str,
    table_index: int,
    row_index: int,
    col_index: int,
    top: float = 0,
    bottom: float = 0,
    left: float = 0,
    right: float = 0
) -> Dict[str, Any]:
    """
    设置单元格内边距

    参数:
        filename: 文档路径
        table_index: 表格索引
        row_index: 行索引
        col_index: 列索引
        top/bottom/left/right: 内边距（英寸）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]
    cell = table.cell(row_index, col_index)

    # 设置内边距（通过XML）
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for side, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if value > 0:
            mar = OxmlElement(f'w:{side}')
            mar.set(qn('w:w'), str(int(value * 1440)))  # 转换为twips
            mar.set(qn('w:type'), 'dxa')
            tcMar.append(mar)

    tcPr.append(tcMar)
    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"单元格内边距设置成功"
    }


@handle_docx_errors
async def set_row_height(
    filename: str,
    table_index: int,
    row_index: int,
    height: float
) -> Dict[str, Any]:
    """
    设置表格行高

    参数:
        filename: 文档路径
        table_index: 表格索引
        row_index: 行索引
        height: 行高（英寸）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]
    row = table.rows[row_index]
    row.height = Inches(height)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"行高设置成功（{height}英寸）"
    }


@handle_docx_errors
async def format_cell_text(
    filename: str,
    table_index: int,
    row_index: int,
    col_index: int,
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
    highlight: Optional[str] = None
) -> Dict[str, Any]:
    """
    设置单元格文本格式

    参数:
        filename: 文档路径
        table_index: 表格索引
        row_index: 行索引
        col_index: 列索引
        font_name: 字体名称
        font_size: 字号
        bold: 是否粗体
        italic: 是否斜体
        color: 文字颜色（十六进制）
        highlight: 背景色（高亮），十六进制RGB（可选）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]
    cell = table.cell(row_index, col_index)

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = Pt(font_size)
            if bold is not None:
                run.font.bold = bold
            if italic is not None:
                run.font.italic = italic
            if color:
                run.font.color.rgb = RGBColor(
                    int(color[0:2], 16),
                    int(color[2:4], 16),
                    int(color[4:6], 16)
                )
            if highlight:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), highlight)
                run._element.get_or_add_rPr().append(shd)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"单元格文本格式设置成功"
    }


@handle_docx_errors
async def set_table_indent(
    filename: str,
    table_index: int,
    indent: float
) -> Dict[str, Any]:
    """
    设置表格缩进

    参数:
        filename: 文档路径
        table_index: 表格索引
        indent: 缩进距离（英寸）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]

    # 通过XML设置表格缩进
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(indent * 1440)))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"表格缩进设置成功（{indent}英寸）"
    }


@handle_docx_errors
async def insert_table_column(
    filename: str,
    table_index: int,
    col_index: int
) -> Dict[str, Any]:
    """
    在表格中插入新列

    参数:
        filename: 文档路径
        table_index: 表格索引
        col_index: 插入位置的列索引
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}")

    table = doc.tables[table_index]

    # 在每一行的指定位置插入新单元格（通过XML操作）
    for row in table.rows:
        # 创建新单元格
        new_tc = OxmlElement('w:tc')
        new_tcPr = OxmlElement('w:tcPr')
        new_tcW = OxmlElement('w:tcW')
        new_tcW.set(qn('w:w'), '0')
        new_tcW.set(qn('w:type'), 'auto')
        new_tcPr.append(new_tcW)
        new_tc.append(new_tcPr)

        # 添加段落
        new_p = OxmlElement('w:p')
        new_tc.append(new_p)

        # 插入到指定位置
        row._element.insert(col_index, new_tc)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"表格{table_index}插入新列成功（位置{col_index}）"
    }


@handle_docx_errors
async def get_table_data(
    filename: str,
    table_index: int
) -> Dict[str, Any]:
    """
    读取整个表格的数据

    参数:
        filename: 文档路径
        table_index: 表格索引（从0开始）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}，文档共有{len(doc.tables)}个表格")

    table = doc.tables[table_index]

    # 提取表格数据
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        table_data.append(row_data)

    return {
        "success": True,
        "filename": filename,
        "table_index": table_index,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "data": table_data
    }


@handle_docx_errors
async def get_table_cell_content(
    filename: str,
    table_index: int,
    row_index: int,
    col_index: int
) -> Dict[str, Any]:
    """
    读取指定单元格的内容

    参数:
        filename: 文档路径
        table_index: 表格索引（从0开始）
        row_index: 行索引（从0开始）
        col_index: 列索引（从0开始）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}，文档共有{len(doc.tables)}个表格")

    table = doc.tables[table_index]

    if row_index < 0 or row_index >= len(table.rows):
        raise ValueError(f"行索引超出范围: {row_index}，表格共有{len(table.rows)}行")

    if col_index < 0 or col_index >= len(table.columns):
        raise ValueError(f"列索引超出范围: {col_index}，表格共有{len(table.columns)}列")

    cell = table.rows[row_index].cells[col_index]

    return {
        "success": True,
        "filename": filename,
        "table_index": table_index,
        "row_index": row_index,
        "col_index": col_index,
        "text": cell.text,
        "character_count": len(cell.text)
    }


@handle_docx_errors
async def get_table_info(
    filename: str,
    table_index: int
) -> Dict[str, Any]:
    """
    获取表格的基本信息

    参数:
        filename: 文档路径
        table_index: 表格索引（从0开始）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格索引超出范围: {table_index}，文档共有{len(doc.tables)}个表格")

    table = doc.tables[table_index]

    return {
        "success": True,
        "filename": filename,
        "table_index": table_index,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "style": table.style.name if table.style else ""
    }

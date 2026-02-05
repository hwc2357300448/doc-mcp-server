"""文档基础操作工具"""
import os
from typing import Optional, Dict, Any
from docx import Document
from ..utils import DocumentManager, validate_file_path, handle_docx_errors


# 全局文档管理器实例
doc_manager = DocumentManager()


@handle_docx_errors
async def create_document(
    filename: str,
    title: Optional[str] = None,
    author: Optional[str] = None,
    subject: Optional[str] = None
) -> Dict[str, Any]:
    """
    创建新的Word文档

    参数:
        filename: 文档保存路径
        title: 文档标题（可选）
        author: 作者（可选）
        subject: 主题（可选）
    """
    abs_path = validate_file_path(filename)

    doc = doc_manager.create_new(abs_path)

    # 设置文档属性
    core_properties = doc.core_properties
    if title:
        core_properties.title = title
    if author:
        core_properties.author = author
    if subject:
        core_properties.subject = subject

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"文档创建成功: {filename}",
        "path": abs_path
    }


@handle_docx_errors
async def get_document_info(filename: str) -> Dict[str, Any]:
    """
    获取文档信息和元数据

    参数:
        filename: 文档路径
    """
    abs_path = validate_file_path(filename)

    if not os.path.exists(abs_path):
        raise FileNotFoundError(f"文件不存在: {abs_path}")

    doc = doc_manager.get_or_open(abs_path, reload=True)
    core_props = doc.core_properties

    # 统计信息
    paragraph_count = len(doc.paragraphs)
    table_count = len(doc.tables)

    return {
        "success": True,
        "filename": filename,
        "path": abs_path,
        "title": core_props.title or "",
        "author": core_props.author or "",
        "subject": core_props.subject or "",
        "created": str(core_props.created) if core_props.created else "",
        "modified": str(core_props.modified) if core_props.modified else "",
        "paragraph_count": paragraph_count,
        "table_count": table_count
    }


@handle_docx_errors
async def get_document_text(filename: str) -> Dict[str, Any]:
    """
    提取文档的全部文本内容

    参数:
        filename: 文档路径
    """
    abs_path = validate_file_path(filename)

    if not os.path.exists(abs_path):
        raise FileNotFoundError(f"文件不存在: {abs_path}")

    doc = doc_manager.get_or_open(abs_path, reload=True)

    # 提取所有段落文本
    text_content = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text)

    full_text = "\n".join(text_content)

    return {
        "success": True,
        "filename": filename,
        "text": full_text,
        "paragraph_count": len(text_content),
        "character_count": len(full_text)
    }


@handle_docx_errors
async def list_available_documents(directory: str = ".") -> Dict[str, Any]:
    """
    列出指定目录下的所有Word文档

    参数:
        directory: 目录路径（默认为当前目录）
    """
    abs_dir = os.path.abspath(directory)

    if not os.path.exists(abs_dir):
        raise FileNotFoundError(f"目录不存在: {abs_dir}")

    if not os.path.isdir(abs_dir):
        raise ValueError(f"路径不是目录: {abs_dir}")

    # 查找所有.docx文件
    docx_files = []
    for filename in os.listdir(abs_dir):
        if filename.endswith('.docx') and not filename.startswith('~$'):
            file_path = os.path.join(abs_dir, filename)
            file_stat = os.stat(file_path)
            docx_files.append({
                "filename": filename,
                "path": file_path,
                "size": file_stat.st_size,
                "modified": str(file_stat.st_mtime)
            })

    return {
        "success": True,
        "directory": abs_dir,
        "count": len(docx_files),
        "files": docx_files
    }


@handle_docx_errors
async def copy_document(source_filename: str, destination_filename: Optional[str] = None) -> Dict[str, Any]:
    """
    复制Word文档

    参数:
        source_filename: 源文档路径
        destination_filename: 目标文档路径（可选，默认在同目录下添加_copy后缀）
    """
    import shutil

    source_path = validate_file_path(source_filename)

    if not os.path.exists(source_path):
        raise FileNotFoundError(f"源文件不存在: {source_path}")

    if destination_filename is None:
        base, ext = os.path.splitext(source_path)
        dest_path = f"{base}_copy{ext}"
    else:
        dest_path = validate_file_path(destination_filename)

    shutil.copy2(source_path, dest_path)

    return {
        "success": True,
        "message": "文档复制成功",
        "source": source_path,
        "destination": dest_path
    }


@handle_docx_errors
async def get_paragraph_text(filename: str, paragraph_index: int) -> Dict[str, Any]:
    """
    获取指定段落的文本内容

    参数:
        filename: 文档路径
        paragraph_index: 段落索引（从0开始）
    """
    abs_path = validate_file_path(filename)

    if not os.path.exists(abs_path):
        raise FileNotFoundError(f"文件不存在: {abs_path}")

    doc = doc_manager.get_or_open(abs_path, reload=True)

    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise ValueError(f"段落索引超出范围: {paragraph_index}，文档共有{len(doc.paragraphs)}个段落")

    para = doc.paragraphs[paragraph_index]

    return {
        "success": True,
        "filename": filename,
        "paragraph_index": paragraph_index,
        "text": para.text,
        "style": para.style.name if para.style else "",
        "character_count": len(para.text)
    }


@handle_docx_errors
async def get_paragraph_range_text(
    filename: str,
    start_index: int,
    end_index: int
) -> Dict[str, Any]:
    """
    获取指定范围段落的文本内容（包括表格）

    参数:
        filename: 文档路径
        start_index: 起始段落索引（从0开始，包含）
        end_index: 结束段落索引（从0开始，包含）

    注意：索引是基于文档元素的顺序，包括段落和表格
    """
    abs_path = validate_file_path(filename)

    if not os.path.exists(abs_path):
        raise FileNotFoundError(f"文件不存在: {abs_path}")

    doc = doc_manager.get_or_open(abs_path, reload=True)

    # 获取文档中所有元素的顺序（段落和表格）
    elements = []
    for element in doc.element.body:
        if element.tag.endswith('p'):  # 段落
            para_index = len([e for e in elements if e['type'] == 'paragraph'])
            para = doc.paragraphs[para_index]
            elements.append({
                'type': 'paragraph',
                'content': para
            })
        elif element.tag.endswith('tbl'):  # 表格
            table_index = len([e for e in elements if e['type'] == 'table'])
            table = doc.tables[table_index]
            elements.append({
                'type': 'table',
                'content': table
            })

    total_elements = len(elements)

    if start_index < 0 or end_index >= total_elements:
        raise ValueError(f"元素索引超出范围，文档共有{total_elements}个元素（段落+表格）")

    if start_index > end_index:
        raise ValueError(f"起始索引({start_index})不能大于结束索引({end_index})")

    # 提取范围内的元素
    elements_data = []
    text_parts = []

    for i in range(start_index, end_index + 1):
        element = elements[i]

        if element['type'] == 'paragraph':
            para = element['content']
            elements_data.append({
                "index": i,
                "type": "paragraph",
                "text": para.text,
                "style": para.style.name if para.style else "",
                "character_count": len(para.text)
            })
            text_parts.append(para.text)

        elif element['type'] == 'table':
            table = element['content']
            # 提取表格数据
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)

            # 将表格转换为文本表示
            table_text = "\n".join(["\t".join(row) for row in table_data])

            elements_data.append({
                "index": i,
                "type": "table",
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": table_data,
                "text_representation": table_text
            })
            text_parts.append(f"[表格 {len(table.rows)}行x{len(table.columns)}列]\n{table_text}")

    # 合并所有元素的文本
    combined_text = "\n\n".join(text_parts)

    return {
        "success": True,
        "filename": filename,
        "start_index": start_index,
        "end_index": end_index,
        "element_count": len(elements_data),
        "elements": elements_data,
        "combined_text": combined_text,
        "total_characters": len(combined_text)
    }

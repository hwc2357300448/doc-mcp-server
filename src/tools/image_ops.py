"""图片操作工具"""
import os
from typing import Optional, Dict, Any
from docx.shared import Inches
from ..utils import DocumentManager, validate_file_path, handle_docx_errors

# 全局文档管理器实例
doc_manager = DocumentManager()


@handle_docx_errors
async def insert_image(
    filename: str,
    image_path: str,
    position: Optional[int] = None,
    width: Optional[float] = None,
    height: Optional[float] = None
) -> Dict[str, Any]:
    """
    插入图片到Word文档

    参数:
        filename: 文档路径
        image_path: 图片文件路径
        position: 插入位置（段落索引，从0开始），None表示追加到文档末尾
        width: 图片宽度（英寸，可选）
        height: 图片高度（英寸，可选）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if not os.path.exists(image_path):
        raise FileNotFoundError(f"图片文件不存在: {image_path}")

    # 如果指定了位置，在指定位置插入
    if position is not None:
        if position < 0 or position > len(doc.paragraphs):
            raise ValueError(f"位置索引超出范围: {position}")

        # 在指定位置插入一个新段落
        if position == 0:
            para = doc.paragraphs[0].insert_paragraph_before()
        elif position >= len(doc.paragraphs):
            para = doc.add_paragraph()
        else:
            para = doc.paragraphs[position].insert_paragraph_before()

        # 在新段落中插入图片
        run = para.add_run()
        if width and height:
            run.add_picture(image_path, width=Inches(width), height=Inches(height))
        elif width:
            run.add_picture(image_path, width=Inches(width))
        elif height:
            run.add_picture(image_path, height=Inches(height))
        else:
            run.add_picture(image_path)
    else:
        # 追加到文档末尾
        para = doc.add_paragraph()
        run = para.add_run()
        if width and height:
            run.add_picture(image_path, width=Inches(width), height=Inches(height))
        elif width:
            run.add_picture(image_path, width=Inches(width))
        elif height:
            run.add_picture(image_path, height=Inches(height))
        else:
            run.add_picture(image_path)

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"图片插入成功（位置: {position if position is not None else '文档末尾'}）",
        "image_path": image_path,
        "position": position
    }


@handle_docx_errors
async def delete_image(
    filename: str,
    paragraph_index: int
) -> Dict[str, Any]:
    """
    删除指定段落中的图片

    参数:
        filename: 文档路径
        paragraph_index: 包含图片的段落索引（从0开始）
    """
    abs_path = validate_file_path(filename)
    doc = doc_manager.get_or_open(abs_path)

    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise ValueError(f"段落索引超出范围: {paragraph_index}")

    para = doc.paragraphs[paragraph_index]

    # 检查段落中是否有图片
    has_image = False
    for run in para.runs:
        # 检查 run 中是否包含图片
        if run._element.xpath('.//pic:pic'):
            has_image = True
            # 删除包含图片的 run
            run._element.getparent().remove(run._element)

    if not has_image:
        return {
            "success": False,
            "message": f"段落 {paragraph_index} 中没有找到图片"
        }

    doc_manager.save(abs_path, doc)

    return {
        "success": True,
        "message": f"图片删除成功（段落索引: {paragraph_index}）",
        "paragraph_index": paragraph_index
    }
